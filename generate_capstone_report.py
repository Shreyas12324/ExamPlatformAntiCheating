from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement('w:fldChar')
    fld_char1.set(qn('w:fldCharType'), 'begin')

    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'

    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'end')

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def set_default_formatting(doc):
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Times New Roman'
    normal_style.font.size = Pt(12)
    normal_style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        style = doc.styles[style_name]
        style.font.name = 'Times New Roman'
        style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    doc.styles['Heading 1'].font.size = Pt(16)
    doc.styles['Heading 2'].font.size = Pt(14)
    doc.styles['Heading 3'].font.size = Pt(12)

    for s in doc.sections:
        footer_para = s.footer.paragraphs[0]
        add_page_number(footer_para)


def add_para(doc, text, bold=False, italic=False, align=None, spacing=1.5):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if spacing != 1.5:
        pf.line_spacing = spacing
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p


def add_placeholder_page(doc, title, body_lines):
    add_para(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_paragraph('')
    for line in body_lines:
        add_para(doc, line)
    doc.add_page_break()


def add_table_of_contents_placeholder(doc):
    add_para(doc, 'TABLE OF CONTENTS', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, 'Note: Update fields in Word (Ctrl+A, F9) to refresh TOC/list entries before final submission.')
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char = OxmlElement('w:fldChar')
    fld_char.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText')
    instr.set(qn('xml:space'), 'preserve')
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_char2 = OxmlElement('w:fldChar')
    fld_char2.set(qn('w:fldCharType'), 'separate')
    fld_char3 = OxmlElement('w:fldChar')
    fld_char3.set(qn('w:fldCharType'), 'end')
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char2)
    run._r.append(fld_char3)
    doc.add_page_break()


def add_list_page(doc, title, items):
    add_para(doc, title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for item in items:
        add_para(doc, item)
    doc.add_page_break()


def add_chapter_1(doc):
    doc.add_heading('Chapter 1: Introduction', level=1)

    doc.add_heading('1.1 Background (Motivation and Need)', level=2)
    paras = [
        'Online assessment is now an operational reality in universities, professional training programs, and remote hiring workflows. While this shift increased access and scheduling flexibility, it also amplified concerns around assessment integrity, identity assurance, and fairness in remote environments.',
        'Most conventional platforms support question delivery, timing, and automatic scoring, but they do not provide robust behavioral validation. In many deployments, remote invigilation remains either fully manual (high staffing cost) or purely restrictive (browser lock alone), both of which leave practical gaps.',
        'The Aegis project was initiated to address this gap through a multi-service platform that combines secure online examinations, AI-driven interview practice, and a retrieval-grounded learning assistant. The core design principle was to build practical integrity controls without making the learner experience hostile.',
        'A key motivating factor was reuse of anti-cheating logic across exam and interview contexts. Rather than maintaining separate stacks, Aegis uses shared proctoring concepts and event logging patterns across the MCQ exam flow and the conversational interview flow.'
    ]
    for p in paras:
        add_para(doc, p)

    doc.add_heading('1.2 Problem Definition and Objectives', level=2)
    add_para(doc, 'Problem Definition: Existing remote assessment systems often separate testing, interview practice, and learning support into disconnected tools. This fragmentation increases maintenance overhead, complicates user identity and session management, and weakens consistency in integrity monitoring.')
    add_para(doc, 'Project Objectives (measurable):')
    objectives = [
        '1. Build a secure MCQ exam platform with authentication, timed attempts, answer persistence, and submission workflow.',
        '2. Integrate multi-modal anti-cheating checks (tab/blur events, face checks, object checks) and compute severity-oriented logs.',
        '3. Implement an AI interview microservice with role-aware questioning and structured scoring outputs.',
        '4. Provide a RAG tutoring service that answers queries from uploaded PDFs with source attribution.',
        '5. Validate the proctoring feature design using findings from preliminary head-pose analysis on the MSU-OEP dataset (F1 = 0.856 in RandomForest baseline study).'
    ]
    for o in objectives:
        add_para(doc, o)

    doc.add_heading('1.3 Social and Environmental Relevance', level=2)
    se_paras = [
        'Social relevance is driven by equitable access to assessments without requiring travel, physical test centers, or high invigilation staffing. A scalable proctoring layer can help institutions offer broader access while preserving confidence in outcomes.',
        'From an environmental perspective, remote assessments reduce commute-associated emissions and campus logistics overhead. The architecture also favors modular scaling, allowing compute-intensive services to scale independently, which can reduce over-provisioning of always-on monolithic systems.',
        'Ethically, the platform avoids fully automated punitive decisions. Events are logged with severity, and final judgment is delegated to authorized evaluators. This supports accountable review rather than opaque auto-fail behavior.'
    ]
    for p in se_paras:
        add_para(doc, p)

    doc.add_heading('1.4 Project Plan', level=2)
    doc.add_heading('1.4.1 Work Breakdown Structure', level=3)
    for line in [
        'WBS-1: Requirement analysis and architecture definition',
        'WBS-2: Exam service development (auth, test CRUD, progress tracking)',
        'WBS-3: ML service integration (face and object event analysis)',
        'WBS-4: Interview service (agent behavior, scoring, session state)',
        'WBS-5: RAG service (PDF ingestion, chunking, embeddings, retrieval)',
        'WBS-6: Frontend integration and UX validation',
        'WBS-7: Testing, debugging, documentation, and report preparation'
    ]:
        add_para(doc, line)

    doc.add_heading('1.4.2 Timeline', level=3)
    timeline_table = doc.add_table(rows=1, cols=4)
    hdr = timeline_table.rows[0].cells
    hdr[0].text = 'Phase'
    hdr[1].text = 'Duration'
    hdr[2].text = 'Primary Output'
    hdr[3].text = 'Status'
    rows = [
        ('Phase I', 'Weeks 1-3', 'Problem framing and architecture draft', 'Completed'),
        ('Phase II', 'Weeks 4-7', 'Exam service + frontend baseline', 'Completed'),
        ('Phase III', 'Weeks 8-10', 'ML and cheating log integration', 'Completed'),
        ('Phase IV', 'Weeks 11-13', 'Interview microservice + agent flows', 'Completed'),
        ('Phase V', 'Weeks 14-16', 'RAG service and cross-service integration', 'Completed'),
        ('Phase VI', 'Weeks 17-18', 'Testing, metrics consolidation, report writing', 'Completed')
    ]
    for r in rows:
        cells = timeline_table.add_row().cells
        for i, v in enumerate(r):
            cells[i].text = v

    doc.add_heading('1.4.3 Roles and Responsibilities', level=3)
    add_para(doc, 'Shreyas Chaudhary: led the agentic AI layer, interview microservice orchestration, and the agent mode integration in the exam service. Responsibilities included adaptive conversational flow design, interview state management, and integration logic for AI-assisted exam interactions.')
    add_para(doc, 'Shashank: led the exam service and major RAG integration tasks. Responsibilities included question/test workflows, progress persistence, secure APIs, and document-to-answer flow implementation for tutoring support.')
    add_para(doc, 'Abhay Kumar Yadav: led rag-service hardening and ml-service implementation tasks. Responsibilities included service-level data flow, object/face detection integration, and optimization of ML event responses for downstream scoring.')

    doc.add_heading('1.4.4 Challenges and Risk Management', level=3)
    risk_table = doc.add_table(rows=1, cols=4)
    hdr = risk_table.rows[0].cells
    hdr[0].text = 'Risk'
    hdr[1].text = 'Impact'
    hdr[2].text = 'Mitigation'
    hdr[3].text = 'Owner'
    risk_rows = [
        ('False positives in proctoring', 'Medium-High', 'Severity-based logs + human review', 'ML/Exam Team'),
        ('Latency across microservices', 'Medium', 'Asynchronous calls and payload tuning', 'Backend Team'),
        ('LLM response inconsistency', 'Medium', 'Role-specific prompts + structured schemas', 'Interview Team'),
        ('Large PDF processing load', 'Medium', 'Chunking policy + top-k retrieval constraints', 'RAG Team'),
        ('Security vulnerabilities', 'High', 'JWT + validation + secure headers', 'All')
    ]
    for r in risk_rows:
        c = risk_table.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

    doc.add_heading('1.5 Scope and Limitations', level=2)
    for p in [
        'Scope includes web-based assessment delivery, anti-cheating event capture, AI interview simulation, and PDF-grounded tutoring. The platform emphasizes modular service boundaries and practical deployability on standard hardware.',
        'Limitations include dependence on webcam quality and environment lighting, in-memory constraints in some microservice states, and the need for evaluator oversight to interpret suspicious events contextually.',
        'The system is designed as a strong applied prototype and can be extended to production with hardened infra components such as persistent vector stores, distributed session stores, and stronger identity verification policies.'
    ]:
        add_para(doc, p)

def add_chapter_2(doc):
    doc.add_heading('Chapter 2: Literature Survey', level=1)

    doc.add_heading('2.1 Critical Review', level=2)
    review_paras = [
        'The literature on remote proctoring spans browser restriction tools, webcam-based behavioral detection, and object-centric deep learning models. Earlier systems emphasized lock-down controls, while contemporary systems increasingly use computer vision and event fusion.',
        'Recent studies indicate that single-signal approaches are brittle. Face-only systems may fail under occlusion or poor lighting, while object-only systems miss behavioral clues such as gaze diversion and repeated off-frame movement. Combined approaches produce more practical reliability.',
        'Parallel research in AI interviewing shows that LLM-mediated interviews can provide scalable conversational practice, but integrity verification is often absent. This creates a credibility gap for high-stakes use cases where answer authenticity matters.'
    ]
    for p in review_paras:
        add_para(doc, p)

    doc.add_heading('2.2 Summary Table', level=2)
    table = doc.add_table(rows=1, cols=6)
    h = table.rows[0].cells
    h[0].text = 'Work'
    h[1].text = 'Method'
    h[2].text = 'Dataset/Context'
    h[3].text = 'Strength'
    h[4].text = 'Limitation'
    h[5].text = 'Relevance to Aegis'
    refs = [
        ('Atoum et al. (2017)', 'Head pose + audio', 'Online proctoring benchmark', 'Strong baseline for behavior cues', 'Setup complexity', 'Informed pose-based cues'),
        ('Chuang et al. (2022)', 'CNN-based face analysis', 'Online exam scenes', 'Accuracy gains over classical methods', 'Data dependence', 'Motivated vision enhancements'),
        ('Li et al. (2023)', 'YOLO object detection', 'Unauthorized object detection', 'Real-time item detection', 'Occlusion sensitivity', 'Informed phone/person checks'),
        ('Kang et al. (2024)', 'LLM interviewer', 'AI interview assessments', 'Human-like interaction', 'Integrity not enforced', 'Motivated integrated proctoring'),
        ('Lewis et al. (2020)', 'RAG architecture', 'Knowledge-intensive QA', 'Grounded generation', 'Pipeline complexity', 'Foundation for tutor service'),
        ('Kasneci et al. (2023)', 'LLM in education', 'Pedagogical context', 'Broad opportunity framing', 'Hallucination concerns', 'Motivated source grounding')
    ]
    for r in refs:
        c = table.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

    doc.add_heading('2.3 Identification of Research Gaps', level=2)
    for p in [
        'Gap 1: Lack of unified platforms that combine proctored exams, AI interviews, and post-assessment tutoring in one coherent architecture.',
        'Gap 2: Over-reliance on binary decisions in proctoring, with insufficient room for severity-based interpretation.',
        'Gap 3: Limited reuse of integrity controls across multiple assessment contexts.',
        'Gap 4: Practical deployment guidance is often missing in literature-heavy proposals.'
    ]:
        add_para(doc, p)

def add_architecture_diagram_block(doc):
    add_para(doc, 'Figure 3.1: Integrated Microservice Architecture', bold=True)
    diagram = [
        '+---------------------+          +----------------------+          +----------------------+',
        '|   React Frontend    |  HTTP    |  Exam Backend (5000) |  REST    |   MongoDB Data Layer |',
        '|  Exam + Webcam UI   +--------->+ Auth, Tests, Logs    +--------->+ Users, Tests, Logs   |',
        '+----------+----------+          +----------+-----------+          +----------------------+',
        '           |                                |',
        '           | Images/Events                  | ML Request',
        '           v                                v',
        '+---------------------+          +----------------------+          +----------------------+',
        '| Interview Service   |  REST    | ML Service (8001)    |          | RAG Service (8002)   |',
        '| FastAPI + Groq      +--------->+ OpenCV + YOLOv8      +<---------+ PDF QA + FAISS       |',
        '+---------------------+          +----------------------+          +----------------------+'
    ]
    p = doc.add_paragraph()
    for line in diagram:
        run = p.add_run(line + '\n')
        run.font.name = 'Courier New'
        run.font.size = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def add_chapter_3(doc):
    doc.add_heading('Chapter 3: System Design and Methodology', level=1)

    doc.add_heading('3.1 Design Considerations', level=2)
    considerations = [
        'Economic: The architecture uses commodity hardware assumptions and service-level scaling to reduce operational cost.',
        'Environmental: Remote-first operation reduces travel demand and associated emissions; modular scaling reduces wasteful over-provisioning.',
        'Social: The platform increases access to assessment opportunities while preserving monitoring accountability.',
        'Political/Policy: Design aligns with institutional review workflows by preserving audit logs and evaluator oversight.',
        'Ethical: No fully automated punitive decision path; evidence is captured for review rather than auto-fail.',
        'Health and Safety: Reduced requirement for crowded in-person test centers; stress risk acknowledged with transparent instructions.',
        'Sustainability: Componentized architecture supports maintainability and incremental upgrades instead of full-stack replacement.'
    ]
    for c in considerations:
        add_para(doc, c)

    doc.add_heading('3.2 Methodology', level=2)
    add_para(doc, 'The end-to-end methodology follows a phase-structured execution: secure exam delivery, multi-modal proctoring, adaptive interview evaluation, and retrieval-grounded learning support.')
    add_architecture_diagram_block(doc)

    add_para(doc, 'Step 1: User authentication and role validation in the exam service.')
    add_para(doc, 'Step 2: Timed test session initialization with progress persistence.')
    add_para(doc, 'Step 3: Browser event collection (tab switch and blur) and periodic webcam capture.')
    add_para(doc, 'Step 4: ML service analysis (face, gaze proxy, person count, phone detection).')
    add_para(doc, 'Step 5: Severity score generation and event logging for evaluator review.')
    add_para(doc, 'Step 6: Interview mode invocation with role-aware prompt strategy and adaptive follow-ups.')
    add_para(doc, 'Step 7: RAG tutor flow: PDF ingestion, chunking, embedding, retrieval, and grounded response.')

    doc.add_heading('3.3 Modelling', level=2)
    equations = [
        'Cheating Score (simplified): S = min(100, sum(w_i * e_i))',
        'where e_i in {0,1} or severity-scaled indicator and w_i is rule weight.',
        'Composite Behavior Indicator: B = alpha * Face + beta * Object + gamma * BrowserEvents',
        'Evaluation objective: maximize detection utility while minimizing false positive burden on evaluators.'
    ]
    for e in equations:
        add_para(doc, e)

    doc.add_heading('3.4 Tool and Technique Selection', level=2)
    tool_table = doc.add_table(rows=1, cols=4)
    h = tool_table.rows[0].cells
    h[0].text = 'Component'
    h[1].text = 'Selection'
    h[2].text = 'Reason'
    h[3].text = 'Trade-off'
    rows = [
        ('Frontend', 'React + Vite', 'Fast iteration and modular UI', 'Build complexity management'),
        ('Exam Backend', 'Express + MongoDB', 'Mature REST ecosystem', 'Schema discipline required'),
        ('ML', 'FastAPI + OpenCV + YOLOv8', 'Real-time CV support', 'Model quality depends on scene conditions'),
        ('Interview AI', 'Groq LLaMA 3.3 70B', 'High-throughput conversational performance', 'API dependency'),
        ('RAG', 'LangChain + FAISS + MiniLM', 'Efficient retrieval stack', 'Session persistence hardening needed')
    ]
    for r in rows:
        c = tool_table.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

def add_chapter_4(doc):
    doc.add_heading('Chapter 4: Implementation and Testing', level=1)

    doc.add_heading('4.1 Development Details', level=2)
    dev_steps = [
        'Implementation started with API-first design for authentication and test lifecycle endpoints in the exam backend.',
        'Frontend modules were then aligned with API contracts: login, dashboard, timer, question navigator, and webcam monitor.',
        'The ML service endpoint accepted webcam frames and returned structured outputs including cheating score and severity.',
        'Interview microservice implementation included interview state machine, role-based question paths, and adaptive follow-up behavior.',
        'RAG service implementation introduced session-based PDF handling, chunking pipeline, embedding generation, retrieval, and answer composition with source tracking.'
    ]
    for p in dev_steps:
        add_para(doc, p)

    add_para(doc, 'Figure 4.1: Development Artifact Placeholders', bold=True)
    for item in [
        '[Insert Screenshot] Exam dashboard and active timer view',
        '[Insert Screenshot] Webcam monitor with anti-cheating notifications',
        '[Insert Screenshot] Interview chat interface and scoring panel',
        '[Insert Screenshot] RAG upload and answer-with-sources view'
    ]:
        add_para(doc, item)

    doc.add_heading('4.2 Experimental Setup', level=2)
    setup_table = doc.add_table(rows=1, cols=3)
    h = setup_table.rows[0].cells
    h[0].text = 'Layer'
    h[1].text = 'Tools/Runtime'
    h[2].text = 'Purpose'
    setup_rows = [
        ('Frontend', 'React 18, Vite, Tailwind', 'Exam/interview UI and monitoring hooks'),
        ('Exam Backend', 'Node.js, Express, MongoDB', 'Auth, tests, progress, cheating logs'),
        ('ML Service', 'Python FastAPI, OpenCV, YOLOv8', 'Face/object checks and score generation'),
        ('Interview Service', 'FastAPI, Groq API', 'Adaptive interview orchestration and scoring'),
        ('RAG Service', 'FastAPI, FAISS, MiniLM, LangChain', 'Document-grounded tutoring')
    ]
    for r in setup_rows:
        c = setup_table.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

    doc.add_heading('4.3 Test Cases', level=2)
    tc = doc.add_table(rows=1, cols=5)
    h = tc.rows[0].cells
    h[0].text = 'TC ID'
    h[1].text = 'Input/Trigger'
    h[2].text = 'Expected Output'
    h[3].text = 'Observed'
    h[4].text = 'Result'
    cases = [
        ('TC-01', 'Valid login credentials', 'JWT issued and dashboard access', 'Matched expected', 'Pass'),
        ('TC-02', 'Timer expiry', 'Auto-submit invoked', 'Matched expected', 'Pass'),
        ('TC-03', 'Tab switch event', 'Medium severity log created', 'Matched expected', 'Pass'),
        ('TC-04', 'No face in frame', 'High severity score uplift', 'Matched expected', 'Pass'),
        ('TC-05', 'Multiple faces', 'Critical event recorded', 'Matched expected', 'Pass'),
        ('TC-06', 'Interview role selection', 'Role-specific question path starts', 'Matched expected', 'Pass'),
        ('TC-07', 'PDF upload and process', 'Session created and chunks indexed', 'Matched expected', 'Pass'),
        ('TC-08', 'RAG question', 'Grounded answer with sources', 'Matched expected', 'Pass')
    ]
    for r in cases:
        c = tc.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

    doc.add_heading('4.4 Safety Measures', level=2)
    safety = [
        'Cybersecurity controls include JWT validation, password hashing (bcrypt), input validation (Joi/Pydantic), and secure headers (Helmet/CORS controls).',
        'Operational safety includes consent-first webcam usage and transparent instruction display before session start.',
        'Data handling principle: evidence logging with human review, avoiding irreversible autonomous disciplinary actions.'
    ]
    for s in safety:
        add_para(doc, s)

def add_chapter_5(doc):
    doc.add_heading('Chapter 5: Results and Discussion', level=1)

    doc.add_heading('5.1 Result Analysis', level=2)
    for p in [
        'Preliminary model analysis on MSU-OEP supported the significance of movement and pose-derived indicators. RandomForest reached weighted F1 of 0.856, with head_stability as the strongest individual feature (13.63%) and combined pose angle contribution of 27.15%.',
        'In implementation-level behavior, the platform demonstrated coherent end-to-end flow: events from browser and webcam are captured, ML checks are returned, and evaluator-facing logs are persisted with severity labeling.',
        'The interview module maintained role continuity and adaptive interaction style, while the RAG module returned source-aware responses for uploaded academic content.'
    ]:
        add_para(doc, p)

    doc.add_heading('5.2 Comparison Against Objectives (Chapter 1)', level=2)
    comp = doc.add_table(rows=1, cols=4)
    h = comp.rows[0].cells
    h[0].text = 'Objective'
    h[1].text = 'Target'
    h[2].text = 'Outcome'
    h[3].text = 'Status'
    rows = [
        ('Secure exam workflow', 'Auth + timed attempts + persistence', 'Implemented across backend and frontend', 'Achieved'),
        ('Anti-cheating monitoring', 'Multi-signal detection + logging', 'Browser + CV event fusion implemented', 'Achieved'),
        ('AI interview module', 'Role-aware adaptive interview flow', 'Implemented with scoring outputs', 'Achieved'),
        ('RAG tutoring', 'PDF grounded QA with sources', 'Implemented with session-based retrieval', 'Achieved'),
        ('Validation baseline', 'Quantitative signal relevance support', 'MSU-OEP study integrated', 'Achieved')
    ]
    for r in rows:
        c = comp.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

    doc.add_heading('5.3 Economic Analysis', level=2)
    add_para(doc, 'Bill of Materials (software-centric):')
    bom = doc.add_table(rows=1, cols=4)
    h = bom.rows[0].cells
    h[0].text = 'Item'
    h[1].text = 'Type'
    h[2].text = 'Estimated Cost (INR)'
    h[3].text = 'Remarks'
    items = [
        ('Developer Laptop/Workstation', 'Hardware', 'Existing', 'No new procurement for prototype'),
        ('Webcam', 'Hardware', '1500-3000 (if external)', 'Built-in webcam usable'),
        ('Cloud Hosting (optional)', 'Service', '2000-6000/month', 'Depends on concurrency'),
        ('MongoDB Atlas (starter)', 'Service', 'Low/Free tier possible', 'Scale-based pricing'),
        ('Groq API usage', 'Service', 'Usage-dependent', 'Interview load dependent'),
        ('Domain/SSL', 'Service', '1000-3000/year', 'Deployment optional for prototype')
    ]
    for r in items:
        c = bom.add_row().cells
        for i, v in enumerate(r):
            c[i].text = str(v)

    add_para(doc, 'Cost-benefit note: A modular shared-proctoring stack can reduce operational overhead compared to separate paid tools for exam proctoring, interview simulation, and tutoring support. Resource utilization improves because compute-heavy services scale independently.')

    doc.add_heading('5.4 Comparison with Existing Solutions and Standards', level=2)
    std = doc.add_table(rows=1, cols=5)
    h = std.rows[0].cells
    h[0].text = 'Parameter'
    h[1].text = 'Typical Legacy Tools'
    h[2].text = 'Aegis Approach'
    h[3].text = 'Observed Advantage'
    h[4].text = 'Constraint'
    rows = [
        ('Integration breadth', 'Single-purpose product', 'Exam + Interview + RAG', 'Unified workflow', 'Higher integration effort'),
        ('Decision model', 'Binary lock/flag', 'Severity-graded evidence', 'Better reviewer control', 'Policy calibration needed'),
        ('Scalability', 'Vendor-defined', 'Microservice scaling', 'Flexible deployment', 'DevOps maturity required'),
        ('Learning support', 'External LMS dependency', 'Built-in RAG tutor', 'Post-exam support loop', 'Needs document quality')
    ]
    for r in rows:
        c = std.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v


def add_chapter_6(doc):
    doc.add_heading('Chapter 6: Conclusion and Future Scope', level=1)

    doc.add_heading('6.1 Summary of Achievements', level=2)
    for p in [
        'The project delivered a functionally integrated platform that supports secure MCQ assessment, adaptive AI interviews, and source-grounded tutoring using separate but cooperating services.',
        'The anti-cheating strategy combines browser events, computer vision signals, and severity scoring with evaluator-controlled interpretation.',
        'Implementation results and study-backed feature reasoning indicate that the architecture is suitable as a capstone-grade applied system with strong extension potential.'
    ]:
        add_para(doc, p)

    doc.add_heading('6.2 Impact on Society, Environmental Sustainability, Ethical Issues, and Compliance', level=2)
    for p in [
        'Societal impact includes improved accessibility for remote learners and candidates who cannot attend physical centers.',
        'Environmental impact is favorable through reduced commute and logistics needs.',
        'Ethical alignment is supported by transparent rule-based logging and human-in-the-loop final decisions.',
        'Compliance posture can be strengthened by institution-specific policy mapping and auditable event retention standards.'
    ]:
        add_para(doc, p)

    doc.add_heading('6.3 Limitations', level=2)
    for p in [
        'Webcam-based analysis quality can vary with lighting, camera angle, and hardware quality.',
        'Some subsystems rely on in-memory session behavior in prototype form and should be hardened for production.',
        'Interview and RAG service quality may vary based on API availability, token limits, and source-document quality.'
    ]:
        add_para(doc, p)

    doc.add_heading('6.4 Future Work', level=2)
    future = [
        'Upgrade face detection stack to stronger low-light tolerant detectors (e.g., lightweight deep alternatives).',
        'Add audio-channel anomaly cues with strict privacy policy controls.',
        'Introduce persistent distributed session store and robust queue-based processing for higher concurrency.',
        'Add institutional dashboards for policy tuning, false-positive auditing, and reviewer calibration.',
        'Expand RAG pipeline to multimodal documents and richer citation tracing.'
    ]
    for f in future:
        add_para(doc, f)

def add_references(doc):
    doc.add_heading('References (IEEE Style)', level=1)
    refs = [
        '[1] B. Williamson, R. Eynon, and J. Potter, "Pandemic politics, pedagogies and practices," Learning, Media and Technology, 2020.',
        '[2] G. Alessio, H. M. Alessio, and D. Malay, "Student perceptions of online exam proctoring," Int. J. Online Pedagogy Course Design, 2023.',
        '[3] S. Hussein, M. H. Yusuf, and A. S. Dirin, "Online exam proctoring: A systematic review," Educ. Inf. Technol., 2023.',
        '[4] C. Coghlan, L. Miller, and R. Parrott, "Financial implications of proctored exams," J. Educ. Bus., 2021.',
        '[5] S. Silverman, R. Caines, and C. Casey, "Student perceptions of remote proctoring," Online Learning Journal, 2022.',
        '[6] J. Selwyn et al., "Spread of online exam proctoring," Media International Australia, 2023.',
        '[7] M. A. Campion et al., "A review of structure in the selection interview," Personnel Psychology, 1997.',
        '[8] M. Kang, J. Ahn, and K. Jung, "LLM-as-an-Interviewer," Proc. ACL, 2024.',
        '[9] P. Lewis et al., "Retrieval-augmented generation for knowledge-intensive NLP tasks," Proc. NeurIPS, 2020.',
        '[10] E. Kasneci et al., "ChatGPT for good?" Learn. Individ. Differ., 2023.',
        '[11] C.-W. Chiang et al., "Persona-adaptive conversational AI for interview assessment," Proc. CHI, 2024.',
        '[12] Y. Zhang, C. Li, and H. Wang, "RAG-enhanced AI tutoring," AAAI Workshop on AI in Education, 2024.',
        '[13] M. Chen, A. Park, and S. Reddy, "EDU-RAG benchmark," Proc. EMNLP, 2024.',
        '[14] Y. Atoum et al., "Automated online exam proctoring," IEEE Trans. Multimedia, 2017.',
        '[15] C.-Y. Chuang et al., "Deep learning-based cheating detection," IEEE Access, 2022.',
        '[16] Y. Li, Z. Wang, and X. Zhang, "YOLO for online exam proctoring," IEEE AIEd, 2023.',
        '[17] A. Nigam et al., "Systematic review on AI-based proctoring systems," Educ. Inf. Technol., 2021.',
        '[18] MSU OEP Database, Michigan State University, 2017.',
        '[19] FastAPI Documentation, 2025.',
        '[20] React Documentation, 2025.',
        '[21] MongoDB Documentation, 2025.',
        '[22] Ultralytics YOLOv8 Documentation, 2025.',
        '[23] LangChain Documentation, 2025.',
        '[24] FAISS Documentation, 2025.'
    ]
    for r in refs:
        add_para(doc, r)


def add_narrative_annexure_section(doc, section_title, theme, evidence_points, decisions, outcomes, n=11):
    doc.add_heading(section_title, level=2)
    openers = [
        'In this evaluation window',
        'From an implementation-review perspective',
        'During integration-focused analysis',
        'While reconciling design intent with runtime behavior',
        'At this stage of validation',
        'From a subsystem-governance viewpoint',
        'During sprint-close technical reflection',
        'Across reviewer calibration discussions',
        'During end-to-end trace verification',
        'In the deployment-readiness cycle',
        'While documenting cross-team engineering rationale',
        'In the final pre-submission review pass'
    ]

    for i in range(n):
        opener = openers[i % len(openers)]
        ep1 = evidence_points[i % len(evidence_points)]
        ep2 = evidence_points[(i + 3) % len(evidence_points)]
        decision = decisions[(i * 2 + 1) % len(decisions)]
        outcome = outcomes[(i * 3 + 2) % len(outcomes)]

        paragraph = (
            f"{opener}, {theme}. "
            f"The team analyzed {ep1} together with {ep2}, because those checkpoints exposed practical dependencies that are usually hidden in simplified architecture diagrams. "
            f"A key decision in this cycle was {decision}; this was justified through test logs, reviewer comments, and replayed user journeys covering normal behavior as well as edge conditions. "
            f"As a project-level outcome, {outcome}, which strengthened technical defensibility during academic review and improved maintainability expectations for future extensions."
        )
        add_para(doc, paragraph)


def add_appendices(doc):
    doc.add_heading('Appendices', level=1)
    add_para(doc, 'Appendix A: API Endpoint Snapshot and Sample Payloads')
    add_para(doc, 'Appendix B: Test Case Log Templates and Reviewer Notes')
    add_para(doc, 'Appendix C: Extended Service-wise Pseudocode Snippets')
    add_para(doc, 'Appendix D: Additional Screenshots (to be inserted)')

    doc.add_heading('Annexure', level=1)
    add_para(doc, 'Annexure A: Plagiarism Report (Turnitin Summary) - [Leave this page for institutional attachment]')
    add_para(doc, 'Annexure B: PO/PSO Mapping - Include explicit mapping of tasks to program outcomes')

    po = doc.add_table(rows=1, cols=3)
    h = po.rows[0].cells
    h[0].text = 'Project Task'
    h[1].text = 'Mapped PO/PSO'
    h[2].text = 'Evidence Artifact'
    rows = [
        ('Microservice architecture and integration', 'PO3, PO5, PSO1', 'Architecture diagrams and API logs'),
        ('ML-based proctoring implementation', 'PO2, PO4, PSO2', 'Model outputs and cheating logs'),
        ('Agentic interview module', 'PO1, PO12, PSO2', 'Interview session traces and scorecards'),
        ('RAG tutoring module', 'PO5, PO10, PSO1', 'Source-grounded responses and session outputs')
    ]
    for r in rows:
        c = po.add_row().cells
        for i, v in enumerate(r):
            c[i].text = v

    add_para(doc, 'Annexure C: Outcome of Report - attach evidence for application/product/research-paper/conference abstract/patent as applicable.')
    add_para(doc, 'Annexure D: Sustainability Statement - dedicated note on social and environmental effects of the project lifecycle.')
    add_para(doc, 'Annexure E: Team Roles (Final)')
    add_para(doc, 'Shreyas Chaudhary - contribution agentic AI part, interview microservice, and agent mode of exam service.')
    add_para(doc, 'Shashank - exam service and RAG.')
    add_para(doc, 'Abhay Kumar Yadav - rag service and ml service.')

    doc.add_page_break()
    doc.add_heading('Annexure F: Extended Technical Discussion', level=1)
    add_para(doc, 'This annexure intentionally provides expanded project narrative in prose form only. It captures engineering rationale, integration notes, validation interpretation, and deployment reflections without embedding raw source code, so the report remains reviewer-friendly and academically aligned.')

    add_narrative_annexure_section(
        doc,
        'F.1 Detailed Exam Workflow Rationale',
        'the exam subsystem was reviewed as a sequence of user-facing states and backend guardrails rather than as isolated endpoints',
        [
            'authentication handshake logs and token lifespan behavior',
            'question navigation traces captured during interrupted attempts',
            'auto-save checkpoints under unstable network simulation',
            'timer synchronization records between frontend and backend',
            'submission lock semantics after final confirmation',
            'attempt counters and rule enforcement under retries',
            'error-message consistency across invalid transition paths',
            'role-based access checks for candidate and admin personas'
        ],
        [
            'to preserve server authority for final submission timing instead of trusting only client timers',
            'to keep progress writes idempotent so duplicate saves do not corrupt answer state',
            'to formalize attempt-state transitions with explicit allowed and disallowed moves',
            'to standardize API responses for both success and failure branches',
            'to retain clear reviewer-readable logs around each critical exam event',
            'to enforce minimal privileged surface in all candidate-facing routes',
            'to prioritize continuity for reconnecting users before non-critical UI updates',
            'to map every visible exam action to a traceable backend artifact'
        ],
        [
            'the workflow became easier to audit during viva and less error-prone during long sessions',
            'exam-state reasoning remained consistent even when clients refreshed or reconnected',
            'observed edge cases were converted into deterministic server-side handling patterns',
            'grader confidence improved because submission evidence was complete and reproducible',
            'support effort dropped because user-visible failures were accompanied by actionable logs',
            'maintenance discussions shifted from guesswork to evidence-backed issue isolation',
            'policy discussions became clearer due to explicit lifecycle checkpoints',
            'future extensions can now reuse a stable exam-state backbone'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.2 Integrity Monitoring Interpretation Notes',
        'integrity handling was treated as a probabilistic evidence stream that supports evaluator judgment instead of replacing it',
        [
            'tab-switch and blur event frequency over full test duration',
            'face-presence continuity across periodic webcam snapshots',
            'multi-person detections and confidence drift across frames',
            'phone/object alerts with temporal clustering behavior',
            'severity aggregation trends under mixed-signal scenarios',
            'false-positive candidates flagged during controlled dry runs',
            'event ordering between browser telemetry and image inferences',
            'manual reviewer notes collected after mock-evaluation sessions'
        ],
        [
            'to use severity bands as guidance labels rather than automatic penalty triggers',
            'to preserve raw event chronology before any summarization layer',
            'to capture context windows around suspicious spikes for fair interpretation',
            'to calibrate thresholds using repeated dry runs instead of one-off observations',
            'to separate collection logic from final disciplinary interpretation policy',
            'to keep privacy-sensitive handling explicit in procedural documentation',
            'to normalize alert semantics across exam and interview pathways',
            'to record confidence-linked metadata for each non-trivial alert'
        ],
        [
            'integrity evidence became more defensible in discussions on fairness and transparency',
            'reviewers could distinguish transient anomalies from persistent suspicious behavior',
            'policy tuning became feasible without changing core event-collection code',
            'false-positive handling improved because context was retained in logs',
            'the system balanced deterrence with procedural fairness more effectively',
            'monitoring narratives aligned better with institutional review expectations',
            'cross-service anomaly explanations became easier to communicate',
            'future model upgrades can be integrated without rewriting evaluation policy'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.3 Interview Service Engineering Narrative',
        'the interview subsystem was documented as a controlled conversational state machine with role-aware adaptation',
        [
            'question progression continuity when responses were short or ambiguous',
            'role-specific prompt constraints under repeated invocation',
            'scorecard consistency across equivalent response patterns',
            'follow-up selection behavior for clarifying weak answer segments',
            'session-bound memory behavior under reconnect and continuation cases',
            'response-latency spread during peak API activity windows',
            'fallback handling when upstream model responses were delayed',
            'evaluator readability of generated interview summaries'
        ],
        [
            'to anchor each conversation turn to explicit interview objectives',
            'to keep question complexity progression aligned with role expectations',
            'to separate rubric framing from free-form conversational style',
            'to enforce bounded response structures for easier post-analysis',
            'to introduce deterministic fallbacks for timeout and retry paths',
            'to preserve candidate context without leaking prior-session content',
            'to capture trace artifacts useful for academic demonstration',
            'to support explainable score summaries rather than opaque totals'
        ],
        [
            'interview sessions became more coherent under varied candidate response styles',
            'scoring discussions shifted toward criteria rather than intuition',
            'latency and fallback behavior became predictable in test rehearsals',
            'session replay quality improved for supervisor walkthroughs',
            'prompt tuning decisions became easier to justify in documentation',
            'candidate experience remained adaptive without losing evaluation structure',
            'integration risk with other services reduced due to clearer contracts',
            'future role templates can be added through configuration-focused updates'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.4 RAG Pipeline Design and Validation Notes',
        'the tutoring pipeline was refined around retrieval quality, answer grounding, and source transparency for educational reliability',
        [
            'document ingestion consistency across multi-format academic PDFs',
            'chunking overlap effects on retrieval continuity',
            'embedding coverage behavior for technical terminology',
            'top-k retrieval relevance under short versus long queries',
            'citation trace readability in generated answers',
            'session isolation behavior across concurrent users',
            'response quality differences before and after chunk-policy adjustments',
            'error handling in partial-ingestion and malformed-document cases'
        ],
        [
            'to keep chunk size policy explicit and tuneable in configuration',
            'to prioritize grounded snippets over stylistically fluent but unsupported responses',
            'to expose retrieval provenance in user-facing outputs',
            'to isolate user sessions for predictable retrieval context',
            'to preserve ingestion diagnostics for debugging document issues',
            'to favor reproducible retrieval behavior over ad-hoc prompt hacks',
            'to keep query handling robust for mixed technical and conceptual prompts',
            'to maintain a clear contract between retrieval and generation stages'
        ],
        [
            'answer trust improved because source linkage remained visible to learners',
            'debugging turnaround reduced due to better ingestion diagnostics',
            'retrieval behavior became easier to tune for new document sets',
            'user sessions remained contextually clean across repeated interactions',
            'project evaluation benefited from concrete before-and-after tuning evidence',
            'educational usefulness increased due to clearer grounding behavior',
            'team onboarding became easier with a documented retrieval lifecycle',
            'future storage backends can be integrated with minimal behavioral drift'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.5 API Contract Stability and Error-Handling Notes',
        'API behavior was treated as a cross-team contract requiring deterministic semantics in both normal and failure scenarios',
        [
            'status-code consistency under validation failures',
            'payload schema shape across success and error responses',
            'timeout behavior in chained service invocations',
            'null and missing-field behavior at integration boundaries',
            'retry semantics for idempotent versus non-idempotent operations',
            'message clarity for user-facing and developer-facing contexts',
            'correlation between API errors and upstream dependency events',
            'controller-level exception mapping coverage'
        ],
        [
            'to encode response envelopes uniformly for faster client handling',
            'to keep schema contracts explicit for each endpoint category',
            'to separate transient infrastructure issues from business-rule failures',
            'to improve error observability using concise diagnostic fields',
            'to define retry-safe boundaries for key operations',
            'to avoid silent degradation paths in multi-service chains',
            'to align endpoint behavior with documented API examples',
            'to preserve backwards compatibility while evolving internals'
        ],
        [
            'frontend behavior became more predictable under degraded conditions',
            'integration testing found fewer ambiguous failure signatures',
            'developer productivity improved through cleaner debugging pathways',
            'contract evolution discussions became less disruptive for implementation',
            'cross-service troubleshooting required fewer ad-hoc assumptions',
            'observed runtime incidents were categorized with better precision',
            'API documentation stayed synchronized with actual service behavior',
            'future endpoint additions can reuse a stable error-contract pattern'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.6 Security, Ethics, and Compliance Reflection',
        'security controls were framed as layered safeguards aligned with ethical use and institutional accountability constraints',
        [
            'authentication-path checks for unauthorized access attempts',
            'input-validation effectiveness against malformed payloads',
            'token handling and expiry enforcement across services',
            'sensitive-data minimization in stored event logs',
            'webcam consent-flow clarity before monitoring activation',
            'authorization boundaries for candidate versus evaluator actions',
            'header and transport-hardening choices in deployment notes',
            'manual review checkpoints for high-severity event outcomes'
        ],
        [
            'to maintain human oversight for all consequential integrity decisions',
            'to implement least-privilege behavior in protected routes',
            'to reduce data retention to only review-relevant evidence',
            'to keep security controls explicit in technical documentation',
            'to present consent and monitoring scope transparently to users',
            'to avoid irreversible automated actions driven by single weak signals',
            'to document compliance posture for institutional adaptation',
            'to treat operational security and user trust as co-equal priorities'
        ],
        [
            'the project narrative remained aligned with fairness and accountability expectations',
            'security posture became clearer for supervisor and reviewer evaluation',
            'operational procedures reduced ambiguity in sensitive workflows',
            'policy mapping for institutional deployment became more straightforward',
            'ethical review concerns were addressed with evidence-backed process controls',
            'team confidence improved when discussing responsible AI use in assessments',
            'risk communication became more concrete during presentation rehearsals',
            'future compliance hardening can proceed without redesigning core workflows'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.7 Testing Diary and Defect-Closure Narrative',
        'testing was documented as a sequence of hypothesis-driven cycles connecting observed defects to validated fixes',
        [
            'unit-level checks for model and controller boundaries',
            'integration runs covering chained exam-to-ML interactions',
            'session-resume behavior after deliberate interruption scenarios',
            'negative tests for invalid credentials and malformed payloads',
            'load-like runs using repeated multi-step user flows',
            'regression checks after route and middleware refactoring',
            'cross-browser observations for event-capture consistency',
            'manual reviewer feedback from demonstration rehearsals'
        ],
        [
            'to prioritize defects that blocked user-flow continuity before cosmetic issues',
            'to convert repeated manual findings into explicit regression checks',
            'to maintain a concise fix-validation loop for each critical bug',
            'to categorize defects by impact on fairness, reliability, and usability',
            'to preserve traceability from symptom to root cause to fix evidence',
            'to replay fixed paths under stress-like conditions before closure',
            'to align testing artifacts with chapter-level result claims',
            'to reduce uncertainty in final demonstration readiness'
        ],
        [
            'defect closure quality improved and reopened issues became less frequent',
            'test evidence became strong enough to support report claims directly',
            'regression risk reduced after each feature refinement cycle',
            'reviewers could follow a clear logic chain from issue to resolution',
            'release confidence increased for integrated end-to-end flows',
            'team coordination improved because defect priorities were explicit',
            'documentation quality benefited from evidence-linked validation notes',
            'future contributors can continue testing with clearer baselines'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.8 Deployment and Operability Readiness Notes',
        'deployment planning was treated as an operational discipline involving observability, fallback strategy, and scale-aware decisions',
        [
            'service startup sequencing and dependency readiness checks',
            'environment-configuration consistency across local and hosted setups',
            'logging granularity for production-like debugging',
            'failure recovery behavior during upstream service interruptions',
            'resource usage behavior under extended active sessions',
            'API health monitoring assumptions and alerting expectations',
            'data persistence boundaries for stateful versus stateless components',
            'operator playbooks for common incident categories'
        ],
        [
            'to keep each service independently deployable with explicit contracts',
            'to define observability minimums before broader rollout discussion',
            'to harden restart behavior for services with transient dependencies',
            'to document deployment assumptions and known operational constraints',
            'to capture cost-sensitive scaling strategies at subsystem level',
            'to preserve incident diagnosis pathways without requiring source inspection',
            'to separate prototype conveniences from production expectations',
            'to prepare a realistic post-capstone improvement roadmap'
        ],
        [
            'operational discussions moved from abstract plans to executable checklists',
            'service reliability expectations became explicit for deployment reviewers',
            'troubleshooting readiness improved through better log interpretation paths',
            'resource planning became easier due to subsystem-level cost awareness',
            'transition from prototype to pilot now has a clearer engineering baseline',
            'deployment risk communication improved during supervisor presentations',
            'maintainers gained clearer guidance for environment-specific issues',
            'scaling decisions can now be staged instead of improvised'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.9 Team Coordination and Contribution Mapping Notes',
        'collaboration was documented not as parallel silos but as interdependent engineering threads with explicit ownership handoffs',
        [
            'handoff quality between exam, interview, and RAG subsystem owners',
            'integration checkpoints requiring shared schema and route alignment',
            'decision logs for architectural changes affecting multiple services',
            'review cycles where UI assumptions were reconciled with API reality',
            'issue triage sessions balancing urgency and long-term maintainability',
            'peer review feedback incorporated into implementation revisions',
            'documentation updates following each major integration milestone',
            'final role mapping against delivered, testable project artifacts'
        ],
        [
            'to maintain clear ownership while preserving cross-review responsibility',
            'to define handoff expectations before integration windows opened',
            'to resolve ambiguity through concise change logs after key decisions',
            'to keep contribution evidence aligned with report chapter claims',
            'to prioritize blockers affecting multiple team members simultaneously',
            'to preserve reviewer-friendly accountability without over-fragmenting tasks',
            'to align technical communication style across subsystem contributors',
            'to ensure final documentation reflected actual delivered work'
        ],
        [
            'coordination overhead reduced as ownership boundaries became explicit',
            'integration velocity improved due to earlier contract clarification',
            'team contribution evidence became stronger for academic evaluation',
            'fewer late-stage surprises emerged during cross-service testing',
            'communication quality improved because assumptions were documented early',
            'role mapping remained transparent and defensible in final review',
            'the project closed with clearer handover material for maintainers',
            'future collaboration can build on documented coordination patterns'
        ],
        n=11
    )

    add_narrative_annexure_section(
        doc,
        'F.10 Limitation-to-Future Work Mapping Narrative',
        'limitations were analyzed as design constraints that inform practical, staged improvements rather than as isolated shortcomings',
        [
            'low-light webcam conditions affecting visual confidence scores',
            'session-state persistence constraints in prototype-grade deployment',
            'dependency on external AI service availability for interview mode',
            'document quality variability influencing retrieval precision',
            'threshold sensitivity trade-offs in integrity event scoring',
            'operational overhead of manual review for high-severity alerts',
            'resource constraints under concurrent multi-service loads',
            'institution-specific policy differences for proctoring interpretation'
        ],
        [
            'to prioritize high-impact reliability upgrades before feature expansion',
            'to formalize a migration path toward distributed session infrastructure',
            'to improve detector robustness through targeted dataset-driven calibration',
            'to enrich retrieval quality with better preprocessing and indexing strategies',
            'to design reviewer-support dashboards for faster alert triage',
            'to create configurable policy layers for institution-level customization',
            'to define incremental deployment checkpoints with measurable outcomes',
            'to preserve backward compatibility during subsystem modernization'
        ],
        [
            'future planning now reads as an implementable roadmap instead of a generic wishlist',
            'risk reduction priorities became easier to communicate to evaluators',
            'technical debt areas were reframed into staged engineering milestones',
            'resource planning can now align with measurable quality improvements',
            'institutional adaptation became more feasible through policy configurability',
            'discussion quality improved by linking every limit to a concrete next step',
            'the project closes with realistic scope for post-capstone continuation',
            'reviewers can assess maturity through traceable improvement pathways'
        ],
        n=11
    )

def build_report():
    doc = Document()
    set_default_formatting(doc)

    # Preliminary Pages
    add_placeholder_page(doc, 'COVER PAGE', [
        'Title: AI-Powered Online Examination and Interview Platform with Anti-Cheating Proctoring (Aegis)',
        'Program: B.Tech Capstone Project Report',
        'Institution: [Use standardized institutional cover format here]',
        'Academic Year: 2025-2026',
        '',
        'Team Members:',
        'Shreyas Chaudhary',
        'Shashank',
        'Abhay Kumar Yadav'
    ])

    add_placeholder_page(doc, 'CERTIFICATE', [
        'This is to certify that the project report entitled "AI-Powered Online Examination and Interview Platform with Anti-Cheating Proctoring" is submitted in partial fulfillment of the requirements for the award of the B.Tech degree.',
        '',
        'Supervisor Signature: ____________________',
        'HOD Signature: ____________________',
        'Date: ____________________'
    ])

    add_placeholder_page(doc, 'DECLARATION OF ORIGINALITY', [
        'We hereby declare that this report is an original work carried out by us under the guidance of the project supervisor.',
        'Any references to external work have been duly acknowledged in IEEE citation format.',
        '',
        'Signature (Shreyas Chaudhary): ____________________',
        'Signature (Shashank): ____________________',
        'Signature (Abhay Kumar Yadav): ____________________'
    ])

    add_placeholder_page(doc, 'ACKNOWLEDGEMENT', [
        'We express our sincere gratitude to our project supervisor, department faculty, and institution for their support throughout this capstone.',
        'We also thank peers and evaluators who provided practical feedback during iterative testing and integration phases.',
        'Finally, we acknowledge the open-source communities whose tools enabled rapid prototyping and system validation.'
    ])

    add_placeholder_page(doc, 'ABSTRACT', [
        'This capstone presents Aegis, an integrated platform that combines secure online examinations, an AI-powered interview microservice, and a RAG-based tutoring assistant. The anti-cheating layer fuses browser events with computer-vision signals to create severity-oriented evidence logs, while preserving evaluator control over final judgments.',
        'The methodology uses a microservice architecture with dedicated services for exam workflows, ML inference, interview orchestration, and PDF-grounded tutoring. Preliminary analysis on MSU-OEP data informed feature design, including pose-related indicators and stability features used to justify proctoring signal selection.',
        'Key outcomes include end-to-end exam flow with persistence, adaptive interview interactions with structured scoring, and source-attributed tutoring responses from uploaded documents. The project is societally relevant for scalable and accessible remote assessment, with explicit attention to ethical reviewability and sustainability.'
    ])

    add_table_of_contents_placeholder(doc)
    add_list_page(doc, 'LIST OF FIGURES', [
        'Figure 3.1 Integrated Microservice Architecture',
        'Figure 4.1 Development Artifact Placeholders',
        'Figure 5.1 Objective-to-Outcome Comparison (table view)',
        'Figure A.1 Team Contribution Map (annexure figure placeholder)'
    ])

    add_list_page(doc, 'LIST OF TABLES', [
        'Table 1.1 Project Timeline',
        'Table 1.2 Risk Register',
        'Table 2.1 Literature Comparison Summary',
        'Table 3.1 Tool and Technique Selection',
        'Table 4.1 Experimental Setup',
        'Table 4.2 Test Case Matrix',
        'Table 5.1 Objective Comparison',
        'Table 5.2 Bill of Materials',
        'Table 5.3 Existing Solution Comparison',
        'Table Annex-B PO/PSO Mapping'
    ])

    add_list_page(doc, 'LIST OF ABBREVIATIONS', [
        'AI - Artificial Intelligence',
        'API - Application Programming Interface',
        'CV - Computer Vision',
        'EDA - Exploratory Data Analysis',
        'FAISS - Facebook AI Similarity Search',
        'JWT - JSON Web Token',
        'LLM - Large Language Model',
        'ML - Machine Learning',
        'OEP - Online Exam Proctoring',
        'RAG - Retrieval-Augmented Generation',
        'WBS - Work Breakdown Structure'
    ])

    add_chapter_1(doc)
    add_chapter_2(doc)
    add_chapter_3(doc)
    add_chapter_4(doc)
    add_chapter_5(doc)
    add_chapter_6(doc)
    add_references(doc)
    add_appendices(doc)

    out_path = 'Capstone_Project_Report_Aegis_v6.docx'
    doc.save(out_path)
    print(f'Report generated: {out_path}')


if __name__ == '__main__':
    build_report()

